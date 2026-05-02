"""
Generates the System Design Document for ResearchAI as a .docx file.
Run:  python3 generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg="1F3864", header_fg="FFFFFF"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(header_fg)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)

    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        bg   = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def page_break(doc):
    doc.add_page_break()

# ── Document ───────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── Cover page ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("System Design Document")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor.from_string("1F3864")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ResearchAI — Intelligent Research Assistant")
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor.from_string("2E75B6")

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %d, %Y')}")

    doc.add_paragraph()
    paradigm = doc.add_paragraph()
    paradigm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = paradigm.add_run("Based on the ReAct (Reasoning + Acting) Paradigm")
    pr.italic = True
    pr.font.size = Pt(12)
    pr.font.color.rgb = RGBColor.from_string("7030A0")

    doc.add_paragraph()

    # Architecture overview table (cover)
    add_table(doc,
        ["Layer", "Component", "Technology Used"],
        [
            ["Reasoning",       "LLM Agent",      "Groq / Llama-3.1-8b-instant"],
            ["Execution",       "Skills / Tools",  "Python Functions, DuckDuckGo API, MySQL"],
            ["Logic Flow",      "ReAct Orchestrator", "Custom ReAct loop (Thought→Action→Observation)"],
            ["Persistence",     "Memory / Storage", "MySQL (structured) + Search History"],
            ["Interoperability","Tool Protocol",   "JSON-based tool registry (MCP-inspired)"],
        ],
        col_widths=[1.5, 1.8, 3.2],
    )

    page_break(doc)

    # ── 1. System Overview ─────────────────────────────────────────────────────
    add_heading(doc, "1. System Overview", level=1)
    add_para(doc,
        "ResearchAI is a web-based intelligent research assistant that enables users to discover, "
        "analyse, and manage academic papers and researchers through natural language. The system "
        "implements the ReAct (Reasoning + Acting) paradigm: the LLM agent reasons about the user's "
        "intent, selects appropriate tools (web search, local database queries, write operations), "
        "observes the results, and iterates until it can produce a final, human-readable answer."
    )
    doc.add_paragraph()
    add_para(doc, "Key capabilities:", bold=True)
    for cap in [
        "Natural-language queries: 'List all papers by Yann LeCun on deep learning'",
        "Real-time web search for papers and researchers via DuckDuckGo",
        "Local MySQL database for persistent paper and researcher storage",
        "Save discovered papers directly to the local library with one command",
        "Full audit trail: every query and response is logged in search_history",
        "Transparent ReAct reasoning — users can inspect the agent's thought process",
    ]:
        add_bullet(doc, cap)

    page_break(doc)

    # ── 2. Function Design ─────────────────────────────────────────────────────
    add_heading(doc, "2. Function Design", level=1)

    add_heading(doc, "2.1 Agent Tools (Skills)", level=2)
    add_para(doc,
        "Each tool is a Python function registered in the TOOLS dictionary. The agent selects "
        "tools by name, passes a plain-text or JSON input, and receives a string observation."
    )
    doc.add_paragraph()
    add_table(doc,
        ["Tool Name", "Input Type", "Description", "Returns"],
        [
            ["search_web",        "Plain text query",  "DuckDuckGo web search for academic content",         "Up to 5 search results with URL"],
            ["search_local_db",   "Plain text keyword","Full-text search across papers and researchers table","Matching rows from MySQL"],
            ["save_paper",        "JSON object",       "Insert a new paper record into the papers table",     "Success/duplicate message"],
            ["list_papers",       "Optional filter",   "List all papers, optionally filtered by year/author", "Formatted paper list"],
            ["list_researchers",  "Optional filter",   "List all researchers, optionally filtered",           "Formatted researcher list"],
            ["add_researcher",    "JSON object",       "Insert a new researcher record into the DB",          "Success/duplicate message"],
        ],
        col_widths=[1.5, 1.3, 2.8, 1.9],
    )

    doc.add_paragraph()
    add_heading(doc, "2.2 ReAct Agent Logic", level=2)
    add_para(doc,
        "The run_agent() function implements the ReAct loop. It sends the user query to the LLM "
        "with a system prompt that describes all available tools and the required output format. "
        "The loop runs for up to 8 iterations:"
    )
    for step in [
        "LLM produces: Thought → Action → Action Input",
        "System executes the named tool with the given input",
        "Observation is appended to the conversation",
        "LLM reads the Observation and decides next step",
        "If LLM outputs 'Final Answer:', the loop terminates",
    ]:
        add_bullet(doc, step)

    doc.add_paragraph()
    add_heading(doc, "2.3 Flask API Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET",  "/",                "Serves the main chat UI (index.html)"],
            ["POST", "/chat",            "Accepts {message}, runs ReAct agent, returns {answer, thoughts}"],
            ["GET",  "/api/stats",       "Returns counts: papers, researchers, search queries"],
            ["GET",  "/api/papers",      "Returns latest 20 papers from MySQL"],
            ["GET",  "/api/researchers", "Returns all researchers from MySQL"],
        ],
        col_widths=[0.8, 2.0, 4.0],
    )

    page_break(doc)

    # ── 3. Main UI Design ──────────────────────────────────────────────────────
    add_heading(doc, "3. Main UI Design", level=1)
    add_para(doc,
        "The frontend is a single-page application served by Flask using plain HTML5, CSS3, "
        "and vanilla JavaScript — no frontend framework required. The layout uses CSS Flexbox "
        "and a dark professional colour scheme."
    )
    doc.add_paragraph()

    add_heading(doc, "3.1 Layout Structure", level=2)
    add_table(doc,
        ["Region", "Width", "Contents"],
        [
            ["Left Sidebar",  "280 px fixed", "Logo, DB stats (papers / researchers / queries), example queries, recent papers list"],
            ["Main Chat Area","Flexible (flex:1)", "Chat header with badge, scrollable message history, ReAct thinking indicator, text input"],
            ["Thought Drawer","560 px slide-in overlay", "Full ReAct loop transcript (Thought / Action / Observation) for each response"],
        ],
        col_widths=[1.5, 1.8, 3.7],
    )

    doc.add_paragraph()
    add_heading(doc, "3.2 Key UI Components", level=2)
    for comp in [
        "Message bubbles: user messages (right-aligned, blue tint) and AI responses (left-aligned, dark surface)",
        "Thinking animation: three bouncing dots shown while the agent is running",
        "Stat cards: live counters for papers, researchers, and total queries",
        "Example query chips: clickable shortcuts that pre-fill the input box",
        "ReAct drawer: slide-in panel showing the agent's full reasoning chain — opened via 'View agent reasoning' link",
        "Auto-resize textarea: input box grows with content, max 160 px",
    ]:
        add_bullet(doc, comp)

    doc.add_paragraph()
    add_heading(doc, "3.3 Colour Palette", level=2)
    add_table(doc,
        ["Variable", "Hex Value", "Usage"],
        [
            ["--bg",        "#0F1117", "Page background"],
            ["--sidebar",   "#161B27", "Sidebar and header"],
            ["--surface",   "#1E2535", "Cards, input, AI bubbles"],
            ["--accent",    "#4F8EF7", "Primary blue — buttons, links, stat numbers"],
            ["--accent2",   "#7C3AED", "Purple — badge gradient"],
            ["--text",      "#E2E8F0", "Primary text"],
            ["--muted",     "#8892A4", "Secondary text, placeholders"],
        ],
        col_widths=[1.5, 1.5, 4.0],
    )

    page_break(doc)

    # ── 4. Main Workflow Design ────────────────────────────────────────────────
    add_heading(doc, "4. Main Workflow Design", level=1)

    add_heading(doc, "4.1 End-to-End Request Flow", level=2)
    add_para(doc, "The following steps describe a complete request from user input to final response:")
    doc.add_paragraph()

    steps = [
        ("User Input",         "User types a natural-language query into the chat UI and presses Enter or Send."),
        ("HTTP POST /chat",    "The browser sends a JSON request {\"message\": \"...\"} to the Flask backend."),
        ("Agent Invocation",   "run_agent() is called. The system prompt (with tool descriptions) and the user query are assembled into a message list."),
        ("LLM Reasoning",      "The Groq API (Llama-3.1-8b) receives the messages and returns a Thought + Action + Action Input block."),
        ("Tool Execution",     "The orchestrator parses the Action name and Action Input, then calls the corresponding Python tool function."),
        ("Observation",        "The tool's return value (string) is appended as an Observation message and fed back to the LLM."),
        ("Iteration",          "Steps 4–6 repeat until the LLM outputs 'Final Answer:' or 8 iterations are reached."),
        ("Response",           "The Flask endpoint returns {answer, thoughts} as JSON. The UI renders the answer and stores the ReAct log."),
        ("DB Logging",         "The query and final answer are inserted into search_history for audit purposes."),
        ("UI Update",          "The sidebar stats and recent papers list refresh via /api/stats and /api/papers."),
    ]

    add_table(doc,
        ["Step", "Actor", "Action"],
        [(f"{i+1}. {s}", a, d) for i, (s, a, d) in
         enumerate([(s.split()[0], s, d) for s, d in steps])],
        col_widths=[1.2, 0, 0],
    )

    # Re-build with correct columns
    tbl = doc.tables[-1]
    doc.element.body.remove(tbl._tbl)

    add_table(doc,
        ["Step", "Description"],
        [(f"{i+1}. {s}", d) for i, (s, d) in enumerate(steps)],
        col_widths=[1.6, 5.4],
    )

    doc.add_paragraph()
    add_heading(doc, "4.2 ReAct Loop Detail", level=2)
    add_para(doc,
        "The ReAct (Reasoning + Acting) paradigm interleaves language model reasoning with "
        "real-world tool execution in a structured loop:"
    )
    doc.add_paragraph()

    react_rows = [
        ["Thought",      "The LLM analyses the query and decides what information it needs and which tool to call."],
        ["Action",       "The LLM specifies the exact tool name (e.g., search_web, save_paper)."],
        ["Action Input", "The LLM provides the tool input — plain text or a JSON object."],
        ["Observation",  "The Python orchestrator executes the tool and returns the real result to the LLM."],
        ["Final Answer", "When sufficient information is available, the LLM produces the human-readable response."],
    ]
    add_table(doc,
        ["ReAct Phase", "Description"],
        react_rows,
        col_widths=[1.5, 5.5],
    )

    page_break(doc)

    # ── 5. E-R Diagram & Database Dictionary ──────────────────────────────────
    add_heading(doc, "5. E-R Diagram & Database Dictionary", level=1)

    add_heading(doc, "5.1 Entity-Relationship Description", level=2)
    add_para(doc,
        "The MySQL database 'research_assistant' contains four tables. The relationships are:"
    )
    for rel in [
        "researchers ──< papers  (one researcher can have many papers; researcher_id FK on papers)",
        "papers ──< local_files  (one paper can have many associated local files; paper_id FK on local_files)",
        "search_history is independent — it logs every user query and agent response",
    ]:
        add_bullet(doc, rel)

    doc.add_paragraph()
    add_heading(doc, "5.2 Table: researchers", level=2)
    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",            "INT",          "PK, AUTO_INCREMENT, NOT NULL", "Unique researcher identifier"],
            ["name",          "VARCHAR(255)", "NOT NULL",                     "Full name of the researcher"],
            ["affiliation",   "VARCHAR(500)", "NULL",                         "University or organisation"],
            ["research_area", "VARCHAR(500)", "NULL",                         "Primary research topics"],
            ["email",         "VARCHAR(255)", "NULL",                         "Contact email (optional)"],
            ["created_at",    "TIMESTAMP",    "DEFAULT CURRENT_TIMESTAMP",    "Record creation time"],
        ],
        col_widths=[1.5, 1.3, 2.2, 2.0],
    )

    doc.add_paragraph()
    add_heading(doc, "5.3 Table: papers", level=2)
    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",            "INT",           "PK, AUTO_INCREMENT, NOT NULL", "Unique paper identifier"],
            ["title",         "VARCHAR(1000)", "NOT NULL",                     "Full paper title"],
            ["authors",       "TEXT",          "NULL",                         "Comma-separated author list"],
            ["abstract",      "TEXT",          "NULL",                         "Paper abstract / summary"],
            ["year",          "INT",           "NULL",                         "Publication year"],
            ["venue",         "VARCHAR(500)",  "NULL",                         "Conference or journal name"],
            ["doi",           "VARCHAR(255)",  "NULL",                         "Digital Object Identifier"],
            ["url",           "VARCHAR(1000)", "NULL",                         "Link to paper (arXiv, etc.)"],
            ["researcher_id", "INT",           "FK → researchers(id), NULL",   "Primary researcher of the paper"],
            ["created_at",    "TIMESTAMP",     "DEFAULT CURRENT_TIMESTAMP",    "Record creation time"],
        ],
        col_widths=[1.4, 1.3, 2.1, 2.2],
    )

    doc.add_paragraph()
    add_heading(doc, "5.4 Table: local_files", level=2)
    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",          "INT",           "PK, AUTO_INCREMENT, NOT NULL", "Unique file identifier"],
            ["filename",    "VARCHAR(500)",  "NOT NULL",                     "Original file name"],
            ["filepath",    "VARCHAR(1000)", "NULL",                         "Absolute path on the server"],
            ["description", "TEXT",          "NULL",                         "User-provided description"],
            ["paper_id",    "INT",           "FK → papers(id), NULL",        "Associated paper (optional)"],
            ["created_at",  "TIMESTAMP",     "DEFAULT CURRENT_TIMESTAMP",    "Upload time"],
        ],
        col_widths=[1.4, 1.3, 2.1, 2.2],
    )

    doc.add_paragraph()
    add_heading(doc, "5.5 Table: search_history", level=2)
    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",         "INT",       "PK, AUTO_INCREMENT, NOT NULL", "Unique log entry identifier"],
            ["query",      "TEXT",      "NOT NULL",                     "User's original natural-language query"],
            ["response",   "TEXT",      "NULL",                         "Agent's final answer"],
            ["created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP",    "Timestamp of the query"],
        ],
        col_widths=[1.4, 1.0, 2.1, 2.5],
    )

    page_break(doc)

    # ── 6. Technology Stack Summary ────────────────────────────────────────────
    add_heading(doc, "6. Technology Stack Summary", level=1)
    add_table(doc,
        ["Layer", "Technology", "Role"],
        [
            ["LLM / Reasoning",  "Groq API (Llama-3.1-8b-instant)", "Core reasoning engine for the ReAct agent"],
            ["Orchestration",    "Custom Python ReAct loop",         "Manages Thought→Action→Observation cycle"],
            ["Web Search Tool",  "DuckDuckGo Search (free, no key)", "Live academic paper and researcher search"],
            ["Database Tool",    "MySQL 9.6 + PyMySQL",              "Persistent local paper and researcher storage"],
            ["Backend Framework","Flask 3 + Flask-CORS",             "REST API server and template rendering"],
            ["Frontend",         "HTML5 / CSS3 / Vanilla JS",        "Single-page chat UI, no framework needed"],
            ["Data Seeding",     "Python seed script",               "Pre-loads 5 researchers and 5 landmark papers"],
        ],
        col_widths=[1.8, 2.5, 2.7],
    )

    # ── Save ───────────────────────────────────────────────────────────────────
    out = "/Users/mikemike/Desktop/research_assistant/ResearchAI_Design_Document.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
